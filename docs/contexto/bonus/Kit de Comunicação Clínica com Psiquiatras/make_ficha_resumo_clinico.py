#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Gera a Ficha de Resumo Clinico em DOCX editavel.
"""

from __future__ import annotations

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


HERE = os.path.dirname(__file__)
OUT = os.path.join(HERE, "ficha_resumo_clinico.docx")

GREEN = "3F6058"
LIGHT_GREEN = "F2F8F7"
LINE = "C5D9D4"
ORANGE = "C47850"
INK = RGBColor(26, 46, 42)
MUTED = RGBColor(74, 102, 96)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color=LINE, size="6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=140, bottom=120, end=140):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, size=10, bold=False, italic=False, color=INK):
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_label(cell, label, hint=""):
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(label)
    style_run(r, size=9.5, bold=True)
    if hint:
        p2 = cell.add_paragraph()
        p2.paragraph_format.space_after = Pt(0)
        r2 = p2.add_run(hint)
        style_run(r2, size=8.5, italic=True, color=MUTED)


def add_blank_lines(cell, count=2):
    for _ in range(count):
        p = cell.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("____________________________________________________________")
        style_run(run, size=9, color=MUTED)


def style_box(cell, fill="FFFFFF"):
    set_cell_shading(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def add_full_box(doc, label, hint="", lines=2):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    style_box(cell, "FBFDFD")
    add_label(cell, label, hint)
    add_blank_lines(cell, lines)
    doc.add_paragraph()


def add_three_col_row(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    labels = [
        ("Paciente", "Iniciais"),
        ("Idade", ""),
        ("Data", ""),
    ]
    for idx, (label, hint) in enumerate(labels):
        cell = table.cell(0, idx)
        style_box(cell, "FBFDFD")
        add_label(cell, label, hint)
        add_blank_lines(cell, 1)
    doc.add_paragraph()


def add_impact_row(doc):
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for idx, label in enumerate(("Sono", "Trabalho/estudo", "Vínculos")):
        cell = table.cell(0, idx)
        style_box(cell, "FBFDFD")
        add_label(cell, label)
        add_blank_lines(cell, 1)
    doc.add_paragraph()


def add_urgency(doc):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    style_box(cell, "FBFDFD")
    add_label(cell, "Urgência percebida")
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run("☐ Sem urgência imediata    ☐ Alinhar em breve    ☐ Urgente hoje")
    style_run(r, size=9.5)
    doc.add_paragraph()


def add_note(doc):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    style_box(cell, LIGHT_GREEN)
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(
        "Nota: registre no prontuário o motivo do contato, para quem comunicou, "
        "quando comunicou e o conteúdo essencial."
    )
    style_run(r, size=8.8, color=MUTED)


def build_doc():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(1.7)
    section.right_margin = Cm(1.7)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    styles["Normal"].font.color.rgb = INK

    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = header_p.add_run("Kit de Comunicação Clínica com Psiquiatras | Dra. Tatiana Gontijo")
    style_run(r, size=8.5, bold=True, color=RGBColor(63, 96, 88))

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("FICHA DE RESUMO CLÍNICO")
    style_run(r, size=18, bold=True, color=RGBColor(63, 96, 88))

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(12)
    r = subtitle.add_run("Uso profissional | Programa Lente Clínica")
    style_run(r, size=10.5, italic=True, color=MUTED)

    intro_table = doc.add_table(rows=1, cols=1)
    intro_cell = intro_table.cell(0, 0)
    style_box(intro_cell, "FFF7F0")
    intro_cell.text = ""
    p = intro_cell.paragraphs[0]
    r = p.add_run(
        "Use esta ficha para organizar o que será compartilhado com o psiquiatra. "
        "Inclua apenas informações relevantes para a continuidade do cuidado."
    )
    style_run(r, size=9.5, italic=True)
    doc.add_paragraph()

    add_three_col_row(doc)
    add_full_box(doc, "Motivo do contato", "Por que estou escrevendo agora?", lines=2)
    add_full_box(doc, "Medicação atual conhecida", "Nome, dose e tempo de uso, se o paciente souber informar", lines=2)
    add_full_box(doc, "O que observei em sessão", "Comportamentos, falas, mudança de padrão e impacto funcional", lines=4)
    add_full_box(doc, "Falas literais do paciente", "Use aspas apenas para frases clinicamente relevantes", lines=2)
    add_full_box(doc, "Tempo de evolução do padrão", "Quando começou? Está piorando, melhorando ou oscilando?", lines=2)
    add_impact_row(doc)
    add_urgency(doc)
    add_full_box(doc, "Pedido objetivo ao psiquiatra", "Reavaliação, orientação de conduta, antecipação de consulta ou alinhamento", lines=2)
    add_note(doc)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("Uso exclusivo para fins educativos | Não substitui avaliação clínica individualizada")
    style_run(r, size=8, color=MUTED)

    return doc


if __name__ == "__main__":
    document = build_doc()
    document.save(OUT)
    print(f"DOCX gerado: {OUT}")
